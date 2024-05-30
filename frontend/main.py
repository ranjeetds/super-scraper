from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from scrapy.crawler import CrawlerProcess
from scrapy import Spider, Request
from bs4 import BeautifulSoup
import requests
import openai
from selenium import webdriver
import os
# from langchain import LanguageChain
from docx import Document
import pandas as pd

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

openai.api_key = os.getenv('OPENAI_API_KEY')

class URLRequest(BaseModel):
    url: str

class DynamicContentSpider(Spider):
    name = "dynamic_spider"
    custom_settings = {
        'LOG_ENABLED': False,
    }

    def __init__(self, start_url, *args, **kwargs):
        super(DynamicContentSpider, self).__init__(*args, **kwargs)
        self.start_urls = [start_url]

    def parse(self, response):
        yield {'url': response.url, 'content': response.text}

        # Follow links
        links = response.css('a::attr(href)').getall()
        for link in links:
            if link.startswith('http'):
                yield Request(url=link, callback=self.parse)

def basic_scraper(url):
    response = requests.get(url)
    return BeautifulSoup(response.text, 'html.parser')

def selenium_scraper(url):
    driver = webdriver.Chrome()
    driver.get(url)
    content = driver.page_source
    driver.quit()
    return BeautifulSoup(content, 'html.parser')

def save_to_doc(data, filename):
    doc = Document()
    for entry in data:
        doc.add_heading(entry['url'], level=1)
        doc.add_paragraph(entry['content'])
    doc.save(filename)

def save_to_excel(data, filename):
    df = pd.DataFrame(data)
    df.to_excel(filename, index=False)

@app.post("/scrape")
async def scrape(request: URLRequest, background_tasks: BackgroundTasks):
    url = request.url

    try:
        soup = basic_scraper(url)
        content = soup.prettify()
    except Exception as e:
        try:
            soup = selenium_scraper(url)
            content = soup.prettify()
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    data = [{"url": url, "content": content}]
    background_tasks.add_task(save_to_doc, data, "scraped_data.docx")
    background_tasks.add_task(save_to_excel, data, "scraped_data.xlsx")

    return {"url": url, "content": content}

@app.post("/crawl")
async def crawl(request: URLRequest, background_tasks: BackgroundTasks):
    url = request.url

    process = CrawlerProcess(settings={
        'FEEDS': {
            'output.json': {'format': 'json'},
        },
    })

    process.crawl(DynamicContentSpider, start_url=url)
    process.start()  # the script will block here until the crawling is finished

    with open('output.json') as f:
        crawled_data = f.read()

    background_tasks.add_task(save_to_doc, crawled_data, "crawled_data.docx")
    background_tasks.add_task(save_to_excel, crawled_data, "crawled_data.xlsx")

    return crawled_data

@app.post("/generate_scraper")
async def generate_scraper(request: URLRequest):
    url = request.url

    prompt = f"Generate Python code to scrape data from the following URL: {url}"
    response = openai.Completion.create(engine="text-davinci-004", prompt=prompt, max_tokens=200)
    generated_code = response.choices[0].text

    return {"generated_code": generated_code}

# @app.post("/langchain_task")
# async def langchain_task(request: URLRequest):
#     url = request.url
#     prompt = f"Analyze and provide insights for the content at: {url}"
    
#     chain = LanguageChain()
#     response = chain.run(prompt)
    
#     return {"langchain_response": response}

# import openai

@app.post("/langchain_task")
async def langchain_task(request: URLRequest):
    url = request.url
    prompt = f"Analyze and provide insights for the content at: {url}"
    
    # Set up OpenAI API key (replace 'YOUR_OPENAI_API_KEY' with your actual API key)
    # openai.api_key = 'YOUR_OPENAI_API_KEY'
    
    # Define the prompt and parameters for the completion
    completion_params = {
        "engine": "gpt-3.5-turbo-instruct",
        "prompt": prompt,
        "max_tokens": 200,
    }
    
    # Send the completion request to OpenAI API
    response = openai.Completion.create(**completion_params)
    
    return {"langchain_response": response.choices[0].text}


if __name__ == '__main__':
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

